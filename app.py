"""
Luluty Job Hunting Tool
Run:  python -m streamlit run app.py
"""

import base64, io, json, os, random, re, time
from collections import Counter
from datetime import datetime, timedelta
from urllib.parse import urlparse

import requests
import streamlit as st
import streamlit.components.v1 as components

st.set_page_config(page_title="Luluty Job Hunting Tool", page_icon="🐸", layout="wide")

HERE = os.path.dirname(os.path.abspath(__file__))
CACHE = os.path.join(HERE, "employers.json")
JOBCACHE = os.path.join(HERE, "jobcache.json")
SEENFILE = os.path.join(HERE, "seen_jobs.json")
TRACKFILE = os.path.join(HERE, "applications.json")

UA = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                  "(KHTML, like Gecko) Chrome/126.0 Safari/537.36",
    "Accept-Language": "en-US,en;q=0.9,ar;q=0.8",
    "Accept": "text/html,application/json,*/*",
}
CACHE_HOURS = 12
MIN_DELAY, MAX_DELAY = 1.5, 4.0

HYPE = [
    "Good job Luluty! 💚", "I'm proud of you 🌸", "You're doing well, beautiful ✨",
    "Look at you go 🐸", "Future data analyst incoming 📊",
    "Taibah's finest 🎓", "You got this, habibti 💅", "Keep cooking Luluty 🔥",
    "Your CV is about to eat 😤", "Recruiters aren't ready 💫",
    "Smartest girl in Medina 🧠", "One click closer 💼",
    "SQL queen behaviour 👑", "They'd be lucky to have you 🍀",
    "Certified problem solver 🛠️", "Big Riyadh energy 🏙️",
    "That degree is paying off 📜", "Built different, honestly 💎",
    "Every application is practice 🎯", "Slow progress is still progress 🌱",
    "You're closer than you think 🚀", "Nobody works harder 💪",
    "Confidence looks good on you 😌", "Rejections are redirections ↩️",
    "Your future boss is scrolling 👀", "Databases fear you 🗄️",
    "The offer is coming, be patient ⏳", "Proud of you, always 🤍",
    "Go get it, Luluty 🏆", "You already did the hard part 🌟",
    "Talent + patience = offer 🧮", "Backing you 100% 🫶",
]

# ═══════════════════════════════════════ style
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@400;500;600;700&display=swap');
html, body, [class*="css"] { font-family:'Space Grotesk',sans-serif; }
#MainMenu, footer {visibility:hidden}
.block-container{padding-top:2rem}

.hero{
  background:linear-gradient(135deg,#6EE7B7 0%,#3B82F6 40%,#A855F7 75%,#EC4899 100%);
  background-size:300% 300%; animation:flow 11s ease infinite;
  border-radius:26px; padding:22px 28px; margin-bottom:6px;
  box-shadow:0 12px 34px rgba(168,85,247,.30);
}
@keyframes flow{0%{background-position:0% 50%}50%{background-position:100% 50%}
100%{background-position:0% 50%}}
.hero h1{color:#fff;margin:0;font-size:2.3rem;letter-spacing:-1.4px}
.hero p{color:#F0FDF4;margin:5px 0 0;font-size:.95rem}

.jobcard{
  border:1px solid #E8EAEE;border-radius:18px;padding:16px 18px;margin-bottom:12px;
  background:#fff;color:#111827;box-shadow:0 1px 3px rgba(16,24,40,.05);
  transition:box-shadow .18s ease, transform .18s ease, border-color .18s ease;
  animation:rise .35s cubic-bezier(.22,1,.36,1) both;
}
.jobcard:hover{border-color:#C4B5FD;transform:translateY(-2px);
  box-shadow:0 8px 24px rgba(124,58,237,.12)}
@keyframes rise{from{opacity:0;transform:translateY(10px)}
to{opacity:1;transform:translateY(0)}}

.badge{display:inline-block;padding:3px 11px;border-radius:999px;
  font-size:.71rem;font-weight:600;margin:2px 4px 2px 0;letter-spacing:.1px}
.b-green{background:#DCFCE7;color:#166534}
.b-blue{background:#DBEAFE;color:#1E40AF}
.b-pink{background:#FCE7F3;color:#9D174D}
.b-gray{background:#F3F4F6;color:#4B5563}
.b-amber{background:#FEF3C7;color:#92400E}
.b-red{background:#FEE2E2;color:#991B1B}
.b-new{background:linear-gradient(90deg,#FBBF24,#F59E0B);color:#fff;
  box-shadow:0 2px 6px rgba(245,158,11,.35)}

.fitbar{height:6px;border-radius:999px;background:#F1F1F4;overflow:hidden;
  margin-top:9px}
.fitfill{height:100%;border-radius:999px;
  background:linear-gradient(90deg,#6EE7B7,#3B82F6,#A855F7);
  animation:grow .7s cubic-bezier(.22,1,.36,1) both}
@keyframes grow{from{width:0}}

.stButton>button{border-radius:14px;font-weight:700;border:2px solid #111;
  box-shadow:3px 3px 0 #111;transition:.12s}
.stButton>button:hover{transform:translate(1px,1px);box-shadow:2px 2px 0 #111}
.stTabs [data-baseweb="tab"]{font-weight:600}
div[data-testid="stExpander"]{border-radius:14px;border:1px solid #E8EAEE}
</style>
""", unsafe_allow_html=True)

try:
    _g = base64.b64encode(open(os.path.join(HERE, "frog.gif"), "rb").read()).decode()
    FROG = f'<img src="data:image/gif;base64,{_g}" width="66" style="border-radius:14px">'
except Exception:
    FROG = '<div style="font-size:3rem">🐸</div>'

st.markdown(f"""
<div class="hero"><div style="display:flex;align-items:center;gap:16px">
{FROG}<div><h1>Luluty Job Hunting Tool</h1>
<p>scrape · match · fix the CV · sort some balls 🫧</p></div></div></div>
""", unsafe_allow_html=True)

_hype_js = json.dumps(HYPE)
components.html(f"""
<div id="hb" style="text-align:center;font-family:'Space Grotesk',system-ui;
 font-size:1.24rem;font-weight:700;padding:12px;background:#FFF1F7;
 border:2px dashed #EC4899;border-radius:18px;color:#9D174D;
 transition:opacity .6s ease">{HYPE[0]}</div>
<script>
const H={_hype_js};let seen=[];const el=document.getElementById('hb');
function pick(){{
  let i;
  do {{ i=Math.floor(Math.random()*H.length); }} while(seen.includes(i)&&seen.length<H.length);
  seen.push(i); if(seen.length>8) seen.shift();
  return H[i];
}}
setInterval(()=>{{
  el.style.opacity=0;
  setTimeout(()=>{{el.textContent=pick();el.style.opacity=1;}},600);
}},8000);
</script>
""", height=70)

# ═══════════════════════════════════════ data
SAUDI_CITIES = {
    "Riyadh":["riyadh","الرياض"], "Jeddah":["jeddah","jedda","جدة"],
    "Medina":["medina","madinah","medinah","المدينة"],
    "Mecca":["mecca","makkah","مكة"], "Dammam":["dammam","الدمام"],
    "Khobar":["khobar","الخبر"], "Dhahran":["dhahran","الظهران"],
    "Jubail":["jubail","الجبيل"], "Yanbu":["yanbu","ينبع"],
    "Tabuk":["tabuk","تبوك"], "Abha":["abha","أبها"], "NEOM":["neom","نيوم"],
    "AlUla":["alula","al ula","العلا"], "Qassim":["qassim","buraidah","القصيم"],
    "Hail":["hail","حائل"], "Jazan":["jazan","jizan","جازان"],
    "Al-Ahsa":["ahsa","hofuf","الأحساء"], "Taif":["taif","الطائف"],
}
SAUDI_GENERIC = ["saudi","ksa","k.s.a","السعودية","المملكة"]

CANDIDATE_SLUGS = [
    "tabby","foodics","salla","careem","bayut","qiddiya-investment-company-1",
    "zaintech","neom","redseaglobal","roshn","noon","kitopi","talabat",
    "accenture","deloitte","sap","oracle","cisco","ericsson","nokia","siemens",
    "schneiderelectric","honeywell","ibm","databricks","snowflake","palantir",
    "servicenow","mongodb","elastic","gitlab",
]
PROBES = {
    "lever":"https://api.lever.co/v0/postings/{s}?mode=json",
    "greenhouse":"https://boards-api.greenhouse.io/v1/boards/{s}/jobs",
    "ashby":"https://api.ashbyhq.com/posting-api/job-board/{s}",
    "recruitee":"https://{s}.recruitee.com/api/offers/",
    "workable":"https://apply.workable.com/api/v1/widget/accounts/{s}",
    "smartr":"https://api.smartrecruiters.com/v1/companies/{s}/postings",
}

GOV_ENTITIES = {
    "GACA":"https://careers.gaca.gov.sa/",
    "ZATCA":"https://careers.zatca.gov.sa/?locale=en_GB",
    "SDAIA":"https://careers.sdaia.gov.sa/",
    "GASTAT":"https://careers.stats.gov.sa/en/",
    "SAMA":"https://careers.sama.gov.sa/en/",
    "CMA":"https://careers.cma.gov.sa/",
    "Ministry of Commerce":"https://talent.mc.gov.sa/en/",
    "MCIT":"https://careers.mcit.gov.sa/",
    "Ministry of Sport":"https://careers.mos.gov.sa/",
    "Ministry of Culture":"https://careers.moc.gov.sa/?locale=en_GB",
    "MISA":"https://jobs.misa.gov.sa/",
    "MODON":"https://careers.modon.gov.sa/",
    "NCEC":"https://careers.ncec.gov.sa/",
    "NCVC":"https://careers.ncvc.gov.sa/",
    "NCM":"https://careers.ncm.gov.sa/",
    "Mawani":"https://careers.mawani.gov.sa/",
    "RGA":"https://careers.rga.gov.sa/",
    "GAMI":"https://jobs.gami.gov.sa/?locale=en_US",
    "Awqaf":"https://careers.awqaf.gov.sa/",
    "Saudi Space Agency":"https://ssa.gov.sa/en/careersLandingPage?path=%2Fcareers%2F",
    "SAIP":"https://www.saip.gov.sa/ar/contact-us/careers",
    "SFD":"https://careers.sfd.gov.sa/",
    "SIDF":"https://www.sidf.gov.sa/en/Careers",
    "National Infrastructure Fund":"https://careers.infra.gov.sa/",
    "RCU AlUla":"https://www.rcu.gov.sa/en/contact/careers",
    "Infath":"https://infath.gov.sa/en/careers/",
    "NCGR":"https://ncgr.gov.sa/ar/employment",
    "CHI":"https://careers.chi.gov.sa/",
    "Sharqia Dev Authority":"https://www.sda.gov.sa/join-us?ltr=true",
    "National Events Center":"https://nec.gov.sa/en/join-us",
    "Dulani":"https://www.dulani.gov.sa/en/join-us",
    "GEOSA":"https://geosa.gov.sa/en/EService/pages/recruitment/jobvacancies.aspx",
    "GAFT":"https://gaft.gov.sa/employment-and-training/employment",
    "LCGPA":"https://careers.lcgpa.gov.sa/",
    "Ministry of Health":"https://www.moh.gov.sa/eservices/cards/pages/alljobs.aspx",
    "Pilgrim Experience Program":"https://careers.pep.gov.sa/ar/",
    "NELC":"https://nelc.gov.sa/en/careers",
    "Shared Services Program":"https://ssp.gov.sa/Recruitements.aspx",
    "Environment Fund":"https://www.ef.gov.sa/en/Pages/join-us.aspx",
    "Taqeem":"https://www.taqeem.gov.sa/jobs",
    "NHIC":"https://nhic.gov.sa/Jobs",
    "RCJY":"https://careers.rcjy.gov.sa/en/",
    "Saudi Tourism Authority":"https://sta.gov.sa/en/main-career-page",
    "Public Health Authority":"https://www.pha.gov.sa/en-us/Pages/joinUS.aspx",
    "SFDA":"https://www.sfda.gov.sa/en/career-vision-and-mission",
    "DGA":"https://career.dga.gov.sa/",
    "GAC":"https://gac.mihnati.com/EN/",
}

EMPLOYERS = {
    "Saudi Aramco":"https://careers.aramco.com/",
    "stc group":"https://careers.stc.com.sa/",
    "Maaden":"https://careers.maaden.com/gb/en",
    "SABIC":"https://jobs.sabic.com/",
    "PIF":"https://www.pif.gov.sa/en/careers/",
    "Al Rajhi Bank":"https://careers.alrajhibank.com.sa/",
    "SNB":"https://www.alahli.com.sa/en/pages/about-us/careers",
    "ACWA Power":"https://careers.acwapower.com/job/",
    "NEOM":"https://www.neom.com/en-us/be-part-of-neom/work-at-neom",
    "Elm":"https://career.elm.sa/elm/",
    "Dr Sulaiman Al Habib":"https://talents.hmg.com/?langcode=en",
    "Saudi Electricity":"https://www.se.com.sa/Careers",
    "Riyad Bank":"https://careers.riyadbank.com/",
    "SAB":"https://careers.sab.com/",
    "Alinma Bank":"https://career.alinma.com/",
    "Mobily":"https://www.mobily.com.sa/wps/portal/web/careers",
    "ROSHN":"https://www.roshn.sa/ar/careers",
    "Red Sea Global":"https://careers.theredsea.sa/",
    "Qiddiya":"https://qiddiya.com/careers/",
    "Diriyah Company":"https://www.diriyahcompany.sa/en/careers",
    "Riyadh Air":"https://www.riyadhair.com/en/careers",
    "New Murabba":"https://newmurabba.com/careers",
    "Saudia":"https://careers.saudia.com/",
    "SAMI":"https://sami.jobs.hr.cloud.sap/",
    "SAMI Advanced Electronics":"https://www.aecl.com/en/jobs/",
    "Tadawul Group":"https://careers.tadawulgroup.sa/?lang=en",
    "solutions by stc":"https://solutions.com.sa/careers/",
    "Bahri":"https://www.bahri.sa/en/careers/",
    "Arab National Bank":"https://careers.anb.com.sa/",
    "Banque Saudi Fransi":"https://www.bsfcareers.sa/",
    "Bank Albilad":"https://www.bankalbilad.com.sa/en/about/pages/e-careers.aspx",
    "SAIB":"https://careers.saib.com.sa/en/",
    "Almarai":"https://www.almarai.com/en/careers/",
    "Bupa Arabia":"https://careers.bupa.com.sa/",
    "Tawuniya":"https://www.tawuniya.com/en/careers",
    "Jarir":"https://jobapp.jarir.com/?lang=sa",
    "Matarat Holding":"https://careers.matarat.com.sa/",
    "flynas":"https://career.flynas.com/",
    "SEVEN":"https://careers.seven.sa/?locale=en_GB",
    "Cruise Saudi":"https://career.cruisesaudi.com/",
    "Soudah Development":"https://soudah.sa/en/work-with-us",
    "Alshaya Group":"https://www.alshaya.com/sa/english/careers/",
    "Chalhoub Group":"https://careers.chalhoubgroup.com/jobs",
    "Apparel Group":"https://www.apparelgroup.com/en/careers/",
    "DP World Saudi":"https://www.dpworld.sa/careers",
    "Budget Saudi":"https://www.budgetsaudi.com/en/careers",
    "Dar Al Riyadh":"https://careers.daralriyadh.com/",
    "Nesma & Partners":"https://www.nesmapartners.com/en/careers",
    "AJEX Logistics":"https://www.aj-ex.com/careers",
    "Salehiya Healthcare":"https://salehiya.com/careers/",
    "Tamkeen Technologies":"https://tamkeentech.sa/careers",
    "Tabuk Pharmaceuticals":"https://careers.tabukpharmaceuticals.com/",
    "Saudi German Health":"https://career.saudigermanhealth.com/",
    "Hikma":"https://www.hikma.com/careers/",
    "Coca-Cola Saudi":"https://www.coca-cola.com.sa/careers/",
    "SPIMACO":"https://careers.spimaco.com.sa/",
    "Al-Dawaa Pharmacies":"https://careers.al-dawaa.com/",
    "Mouwasat":"https://www.mouwasat.com/en/careers",
    "Dallah Health":"https://careers.dallahhealth.com/",
    "Panda Retail":"https://careers.panda.com.sa/en/",
    "Al Othaim Markets":"https://recruitment.othaimmarkets.com/",
    "Zahid Group":"https://careers.zahid.com/",
    "alfanar":"https://jobs.alfanar.com/alfanar/go/All-Openings/4442101/?location=saudi+arabia",
    "AtkinsRealis":"https://careers.atkinsrealis.com/en/global-locations/middle-east/saudi-arabia",
    "Accenture":"https://www.accenture.com/sa-en/careers",
    "BCG":"https://careers.bcg.com/global/en/locations/saudi-arabia",
    "EY Saudi":"https://www.ey.com/en_sa/careers",
    "KPMG Saudi":"https://kpmg.com/sa/en/careers.html",
    "SAP":"https://jobs.sap.com/go/Saudi-Arabia/9009201/",
    "Cisco":"https://careers.cisco.com/global/en/saudi-arabia?from=0&s=1",
    "IBM":"https://www.ibm.com/careers/search?field_keyword_05%5B0%5D=Saudi+Arabia",
    "Ericsson":"https://jobs.ericsson.com/careers?location=Saudi+Arabia&page=1",
    "Schneider Electric":"https://careers.se.com/saudi-arabia",
    "Zain":"https://careers.zain.com/",
    "Landmark Group":"https://efhi.fa.em3.oraclecloud.com/hcmUI/CandidateExperience/en/sites/CX_1/jobs",
}

STOPWORDS = set("""
a an the and or of to in for with on at by from as is are was were be been being
i me my we our you your he she it they them this that these those will would can
could should have has had do does did not no so if then than there here what which
who whom when where why how all any both each few more most other some such only own
same too very just don now up out about into over after before under again
year years month months day days work working works experience skills team teams
role responsibilities requirements ability strong good excellent using use used
new job company across within including etc via per able support
""".split())

NOISE = set("""
saudi arabia arabian ksa riyadh jeddah medina medinah mecca makkah dammam khobar
january february march april june july august september october november december
jan feb mar apr jun jul aug sep sept oct nov dec
university college bachelor science degree gpa graduated school institute academy
taibah king abdulaziz faisal saud fahd prince princess
gmail hotmail outlook yahoo email phone mobile address linkedin github
present current date period duration company ltd llc inc corp group holding
""".split())

SKILL_VOCAB = {
    "sql","mysql","postgresql","sql server","ssms","ssis","ssrs","t-sql",
    "power bi","tableau","looker","qlik","excel","vba","dax","power query",
    "python","pandas","numpy","spss","sas",
    "data analysis","data analytics","data validation","data quality","data entry",
    "etl","data warehouse","data modeling","reporting","dashboard","dashboards",
    "kpi","visualization","statistics","forecasting",
    "database","databases","dba","backup","indexing","query optimization",
    "performance tuning","availability groups","stored procedures","execution plans",
    "active directory","exchange server","windows server","linux","powershell",
    "ticketing","service desk","itil","incident","troubleshooting","help desk",
    "networking","dns","dhcp","vpn","firewall",
    "cybersecurity","endpoint security","threat","siem","vulnerability","compliance",
    "erp","sap","oracle","crm","sharepoint","jira",
    "business analysis","requirements","documentation","process improvement",
    "stakeholder","workflow","gap analysis","quality assurance","quality control",
    "audit","inventory","procurement","project management","agile","scrum",
    "coordination","logistics","operations","customer support","training",
}

SENIOR = ["senior","sr.","staff ","principal","lead ","head of","manager",
          "director","vp ","chief","architect","expert","10+","8+ years"," iii"]
JUNIOR = ["junior","graduate","entry","intern","trainee","associate","fresh",
          "0-2","1-2","analyst program","rotational","tamheer","co-op",
          "cooperative","new grad"]

DEGREE_FIELDS = {
    "computer information systems":["information system","computer information",
        "management information","mis","business systems","systems analysis"],
    "computer science":["computer science","software engineering","informatics"],
    "data science":["data science","data analytics","statistics","machine learning"],
    "business":["business administration","management","marketing","finance"],
    "engineering":["electrical engineering","mechanical","civil","industrial",
                   "telecommunication","network engineering"],
    "accounting":["accounting","audit"],
}

WEAK_VERBS = {
    "responsible for":"Led","worked on":"Delivered","helped":"Supported",
    "assisted":"Supported","did":"Executed","made":"Built",
    "was part of":"Contributed to","handled":"Managed","took care of":"Managed",
    "involved in":"Drove","participated in":"Contributed to",
}


# ═══════════════════════════════════════ text utils
def tokens(t):
    ws = re.findall(r"[a-zA-Z][a-zA-Z+#.\-]{1,}", (t or "").lower())
    return [w.strip(".-") for w in ws if w not in STOPWORDS and len(w) > 2]

def strip_html(s):
    s = re.sub(r"<(script|style)[^>]*>.*?</\1>"," ", s or "", flags=re.S|re.I)
    s = re.sub(r"<[^>]+>"," ", s)
    for a,b in [("&amp;","&"),("&nbsp;"," "),("&#39;","'"),("&quot;",'"')]:
        s = s.replace(a,b)
    return re.sub(r"\s+"," ", s).strip()

def _ngrams(text, n=3):
    ws = re.findall(r"[a-zA-Z][a-zA-Z+#.\-]{1,}", (text or "").lower())
    out = []
    for size in range(1, n+1):
        for i in range(len(ws)-size+1):
            out.append(" ".join(ws[i:i+size]))
    return out

def _has(term, hay):
    return re.search(r"(?<![a-z])"+re.escape(term)+r"(?![a-z])", hay) is not None

def build_profile(text, top_n=60):
    ws = tokens(text)
    prof = {t:c for t,c in Counter(ws).most_common(top_n)}
    for t,c in Counter(f"{a} {b}" for a,b in zip(ws,ws[1:])).most_common(top_n//2):
        if c > 1:
            prof[t] = c*2
    return prof

def detect_field(text):
    low = (text or "").lower()
    sc = {f: sum(low.count(k) for k in kws) for f,kws in DEGREE_FIELDS.items()}
    best = max(sc, key=sc.get)
    return (best, DEGREE_FIELDS[best]) if sc[best] > 0 else (None, [])


# ── PDF extraction: try 3 methods, keep the cleanest
def _fused_ratio(txt):
    ws = txt.split()
    if not ws: return 1.0
    return sum(1 for w in ws if len(w) > 25) / len(ws)


def read_upload(f):
    n, data = f.name.lower(), f.read()
    if not n.endswith(".pdf"):
        if n.endswith(".docx"):
            import docx
            return "\n".join(p.text for p in docx.Document(io.BytesIO(data)).paragraphs)
        return data.decode("utf-8", errors="ignore")

    cands = []
    try:                                   # 1 — PyMuPDF, usually best spacing
        import fitz
        doc = fitz.open(stream=data, filetype="pdf")
        cands.append("\n".join(p.get_text(sort=True) for p in doc))
        doc.close()
    except Exception:
        pass
    try:                                   # 2 — pdfplumber, tight tolerance
        import pdfplumber
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            cands.append("\n".join(
                (p.extract_text(x_tolerance=1.2, y_tolerance=3) or "")
                for p in pdf.pages))
    except Exception:
        pass
    try:                                   # 3 — rebuild from word boxes
        import pdfplumber
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            out = []
            for p in pdf.pages:
                lines = {}
                for w in p.extract_words(x_tolerance=1.2):
                    lines.setdefault(round(w["top"]/3), []).append((w["x0"], w["text"]))
                out.append("\n".join(" ".join(t for _,t in sorted(v))
                                     for _,v in sorted(lines.items())))
            cands.append("\n".join(out))
    except Exception:
        pass

    cands = [c for c in cands if c and c.strip()]
    if not cands: return ""
    return min(cands, key=lambda c: (_fused_ratio(c), -len(c)))


SECTION_HEADS = ["education","skills","technical skills","soft skills",
                 "certification","certifications","experience","projects",
                 "summary","professional summary"]

def parse_cv(text):
    lines = [l.rstrip() for l in text.split("\n")]
    cv = {"name":"", "headline":"", "contact":"", "sections":{}}
    for l in lines[:6]:
        s = l.strip()
        if s and s.isupper() and len(s.split()) <= 5 and not cv["name"]:
            cv["name"] = s
        elif s and not cv["headline"] and cv["name"] and len(s) < 70:
            cv["headline"] = s
    if not cv["name"]:
        cv["name"] = next((l.strip() for l in lines if l.strip()), "YOUR NAME")

    email = re.search(r"[\w.+-]+@[\w-]+\.[\w.]+", text)
    phone = re.search(r"(\+?966[\d\s-]{7,}|0?5\d[\d\s-]{6,})", text)
    cv["contact"] = " | ".join(filter(None, [
        email.group(0) if email else "",
        phone.group(0).strip() if phone else "",
        "Saudi Arabia"]))

    cur, buf = "header", []
    for l in lines:
        low = l.strip().lower()
        matched = next((h for h in SECTION_HEADS
                        if low == h or low.startswith(h)), None)
        if matched and len(l.strip()) < 40:
            cv["sections"].setdefault(cur, []).extend(buf)
            cur, buf = matched, []
        else:
            if l.strip(): buf.append(l.strip())
    cv["sections"].setdefault(cur, []).extend(buf)
    return cv


def build_cv_profile(text):
    """Weight by CV section, not word frequency. Skills > experience > certs."""
    cv = parse_cv(text)
    S = cv["sections"]
    prof = {}

    def add(chunk, base):
        if not chunk: return
        for g in _ngrams(chunk):
            if g in SKILL_VOCAB:
                prof[g] = max(prof.get(g,0), base+8)
            elif " " not in g and 4 <= len(g) <= 24 \
                 and g not in NOISE and g not in STOPWORDS:
                prof[g] = max(prof.get(g,0), base)

    for line in (S.get("skills",[]) + S.get("technical skills",[]) +
                 S.get("soft skills",[])):
        for part in re.split(r"[,\u2022•|/]", line):
            p = part.strip().lower()
            if 2 < len(p) < 45 and p not in NOISE and len(p.split()) <= 5:
                prof[p] = 10
                for g in _ngrams(p):
                    if g in SKILL_VOCAB: prof[g] = 12

    add(" ".join(S.get("experience",[]) + S.get("projects",[])), 6)
    add(" ".join(S.get("certification",[]) + S.get("certifications",[])), 4)

    for k in detect_field(text)[1]:
        prof[k] = 9

    return dict(sorted(prof.items(), key=lambda x:-x[1])[:80])


# ═══════════════════════════════════════ persistence
def _load(path, default):
    if os.path.exists(path):
        try: return json.load(open(path, encoding="utf-8"))
        except Exception: pass
    return default

def _save(path, obj):
    try: json.dump(obj, open(path,"w",encoding="utf-8"), ensure_ascii=False)
    except Exception: pass

def job_key(j):
    return f"{j.get('company','')}|{j.get('title','').lower().strip()}"

def load_cache(): return _load(JOBCACHE, {})
def save_cache(c): _save(JOBCACHE, c)

def cache_fresh(e):
    try:
        return datetime.now()-datetime.fromisoformat(e["ts"]) < timedelta(hours=CACHE_HOURS)
    except Exception:
        return False


# ═══════════════════════════════════════ fetching
def polite_sleep(): time.sleep(random.uniform(MIN_DELAY, MAX_DELAY))

def get(url, timeout=20):
    try:
        r = requests.get(url, timeout=timeout, headers=UA)
        if r.status_code == 429:
            time.sleep(60); return None
        return r if r.status_code == 200 else None
    except requests.RequestException:
        return None

def probe(slug):
    for ats,tpl in PROBES.items():
        r = get(tpl.format(s=slug), timeout=10)
        if r:
            try: d = r.json()
            except Exception: continue
            if isinstance(d,list) and d: return ats
            if isinstance(d,dict):
                for k in ("jobs","offers","content","results"):
                    if d.get(k): return ats
        time.sleep(0.4)
    return None

def fetch_ats(slug, ats):
    out = []
    try:
        if ats=="lever":
            r=get(f"https://api.lever.co/v0/postings/{slug}?mode=json")
            for j in (r.json() if r else []):
                out.append({"title":j.get("text",""),
                    "location":(j.get("categories") or {}).get("location",""),
                    "url":j.get("hostedUrl",""),
                    "posted":(j.get("createdAt") and
                              datetime.fromtimestamp(j["createdAt"]/1000).strftime("%Y-%m-%d")) or "",
                    "desc":strip_html(j.get("descriptionPlain") or "")})
        elif ats=="greenhouse":
            r=get(f"https://boards-api.greenhouse.io/v1/boards/{slug}/jobs?content=true")
            for j in (r.json().get("jobs",[]) if r else []):
                out.append({"title":j.get("title",""),
                    "location":(j.get("location") or {}).get("name",""),
                    "url":j.get("absolute_url",""),
                    "posted":(j.get("updated_at") or "")[:10],
                    "desc":strip_html(j.get("content",""))})
        elif ats=="ashby":
            r=get(f"https://api.ashbyhq.com/posting-api/job-board/{slug}")
            for j in (r.json().get("jobs",[]) if r else []):
                out.append({"title":j.get("title",""),"location":j.get("location",""),
                    "url":j.get("jobUrl",""),
                    "posted":(j.get("publishedAt") or "")[:10],
                    "desc":strip_html(j.get("descriptionPlain",""))})
        elif ats=="recruitee":
            r=get(f"https://{slug}.recruitee.com/api/offers/")
            for j in (r.json().get("offers",[]) if r else []):
                out.append({"title":j.get("title",""),"location":j.get("location",""),
                    "url":j.get("careers_url",""),
                    "posted":(j.get("published_at") or "")[:10],
                    "desc":strip_html(j.get("description",""))})
        elif ats=="workable":
            r=get(f"https://apply.workable.com/api/v1/widget/accounts/{slug}?details=true")
            for j in (r.json().get("jobs",[]) if r else []):
                out.append({"title":j.get("title",""),
                    "location":j.get("location") or j.get("city",""),
                    "url":j.get("url",""),
                    "posted":(j.get("published_on") or "")[:10],
                    "desc":strip_html(j.get("description",""))})
        elif ats=="smartr":
            r=get(f"https://api.smartrecruiters.com/v1/companies/{slug}/postings?limit=100")
            for j in (r.json().get("content",[]) if r else []):
                loc=j.get("location") or {}
                out.append({"title":j.get("name",""),
                    "location":f"{loc.get('city','')} {loc.get('country','')}".strip(),
                    "url":j.get("ref") or j.get("applyUrl",""),
                    "posted":(j.get("releasedDate") or "")[:10],"desc":""})
    except Exception:
        pass
    for j in out: j["company"]=slug
    return out

JOB_WORD = re.compile(r"(analyst|engineer|specialist|officer|developer|consultant|"
    r"coordinator|administrator|architect|scientist|technician|associate|intern|"
    r"trainee|graduate|advisor|auditor|designer|researcher)", re.I)

SF_HINT = re.compile(r"locale=en_|/search|successfactors|rmkcdn|hr\.cloud\.sap", re.I)


def fetch_site_browser(entities, progress=None):
    from playwright.sync_api import sync_playwright

    jobs_by_name, dead = {}, []
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        ctx = browser.new_context(user_agent=UA["User-Agent"], locale="en-US",
                                  viewport={"width":1366,"height":900})
        page = ctx.new_page()
        page.route("**/*", lambda r: r.abort()
                   if r.request.resource_type in ("image","font","media","stylesheet")
                   else r.continue_())

        for idx,(name,url) in enumerate(entities):
            found, seen = [], set()
            tries = [url]
            if SF_HINT.search(url) or url.rstrip("/").count("/") <= 2:
                pr = urlparse(url)
                tries.insert(0, f"{pr.scheme}://{pr.netloc}/search/?q=")

            for attempt in tries:
                try:
                    page.goto(attempt, timeout=25000, wait_until="domcontentloaded")
                    try: page.wait_for_load_state("networkidle", timeout=7000)
                    except Exception: pass
                    time.sleep(1.2)
                    page.mouse.wheel(0, 3000)
                    time.sleep(0.8)

                    texts = page.eval_on_selector_all(
                        "a, li, h2, h3, h4, td, .jobTitle, .job-title, "
                        "[class*='job'], [class*='vacancy'], [id*='job']",
                        "els => els.map(e => (e.innerText||'').trim())")

                    for t in texts:
                        t = re.sub(r"\s+"," ", t or "").strip()
                        if (6 < len(t) < 90 and JOB_WORD.search(t)
                                and t.lower() not in seen):
                            seen.add(t.lower())
                            found.append({"title":t,"location":"Saudi Arabia",
                                          "url":page.url,"desc":"","posted":"",
                                          "company":name})
                    if found: break
                except Exception:
                    continue

            if found: jobs_by_name[name] = found[:40]
            else: dead.append(name)
            if progress: progress(idx+1, len(entities), name)
            time.sleep(random.uniform(1.0, 2.2))

        browser.close()
    return jobs_by_name, dead


def fetch_site_plain(name, url):
    r = get(url)
    if not r: return []
    out, seen = [], set()
    for c in re.findall(r"<(?:a|li|h[2-4])[^>]*>(.*?)</(?:a|li|h[2-4])>",
                        r.text, flags=re.S|re.I):
        t = strip_html(c)
        if 6 < len(t) < 90 and JOB_WORD.search(t) and t.lower() not in seen:
            seen.add(t.lower())
            out.append({"title":t,"location":"Saudi Arabia","url":url,
                        "desc":"","posted":"","company":name})
    return out[:40]


def fetch_jobspy(term, cities, all_saudi):
    try:
        from jobspy import scrape_jobs
    except ImportError:
        return [], "JobSpy not installed — run: pip install -U python-jobspy"
    loc = "Saudi Arabia" if all_saudi else f"{cities[0]}, Saudi Arabia"
    try:
        df = scrape_jobs(site_name=["indeed","bayt"], search_term=term,
                         location=loc, results_wanted=25, hours_old=168,
                         country_indeed="Saudi Arabia")
        out = []
        for _,r in df.iterrows():
            posted = r.get("date_posted")
            out.append({"title":str(r.get("title") or ""),
                        "location":str(r.get("location") or ""),
                        "url":str(r.get("job_url") or ""),
                        "posted":str(posted)[:10] if posted and str(posted)!="nan" else "",
                        "desc":str(r.get("description") or "")[:4000],
                        "company":str(r.get("company") or "board")})
        return out, None
    except Exception as e:
        return [], f"JobSpy error: {e}"


# ═══════════════════════════════════════ scoring
def loc_ok(loc, chosen, all_saudi, keep_blank):
    l = (loc or "").lower().strip()
    if not l: return keep_blank
    if all_saudi:
        return (any(g in l for g in SAUDI_GENERIC)
                or any(k in l for c in SAUDI_CITIES.values() for k in c))
    return any(k in l for city in chosen for k in SAUDI_CITIES[city])


def score(job, prof, entry_only, field_kws):
    tl = job["title"].lower()
    desc = (job.get("desc") or "").lower()
    if entry_only and any(m in tl for m in SENIOR): return 0, [], []

    pts, hits = 0, []
    strong = [t for t,w in prof.items() if w >= 8]
    for term, w in prof.items():
        if _has(term, tl):
            pts += w * 4; hits.append(term)
        elif desc and _has(term, desc):
            pts += w
            if w >= 8: hits.append(term)

    if not desc and pts:
        pts = int(pts * 1.8)
    if entry_only and any(m in tl for m in JUNIOR): pts += 25

    hits = sorted(set(hits), key=len, reverse=True)
    gaps = [t for t in strong if t not in hits][:6] if desc else []
    return pts, hits[:8], gaps


@st.cache_resource(show_spinner=False)
def _embedder():
    from sentence_transformers import SentenceTransformer
    return SentenceTransformer("all-MiniLM-L6-v2")


def semantic_rerank(jobs, cv_text, weight=60):
    if not jobs or not cv_text:
        return jobs
    from sentence_transformers import util
    m = _embedder()
    cv_vec = m.encode(cv_text[:4000], convert_to_tensor=True,
                      normalize_embeddings=True)
    texts = [f"{j['title']}. {(j.get('desc') or '')[:900]}" for j in jobs]
    job_vecs = m.encode(texts, convert_to_tensor=True,
                        normalize_embeddings=True, batch_size=32)
    sims = util.cos_sim(cv_vec, job_vecs)[0].tolist()
    for j, s in zip(jobs, sims):
        j["sim"] = round(max(s, 0.0), 3)
        j["score"] = int(j["score"] + j["sim"] * weight)
    return sorted(jobs, key=lambda x: -x["score"])


# ═══════════════════════════════════════ CV tailoring
def extract_bullets(lines):
    out = []
    for l in lines:
        m = re.match(r"^[-•·▪*]\s*(.+)$", l.strip())
        if m: out.append(m.group(1).strip())
        elif len(l.strip()) > 45 and l.strip()[0].isupper(): out.append(l.strip())
    return out

def jd_keywords(jd, n=30):
    ws = tokens(jd)
    uni = Counter(ws)
    bi = Counter(f"{a} {b}" for a,b in zip(ws,ws[1:]))
    out = [t for t,c in bi.most_common(n) if c > 1]
    out += [t for t,c in uni.most_common(n*2) if t not in " ".join(out)]
    return out[:n]

def jd_title(jd):
    m = re.search(r"(?:job title|position|role)\s*[:\-]\s*(.{3,60})", jd, re.I)
    if m: return m.group(1).strip().title()
    for l in jd.split("\n")[:6]:
        s = l.strip()
        if 6 < len(s) < 60 and JOB_WORD.search(s):
            return s.title()
    return "Target Role"

def tailor_bullet(b, kws):
    new = b
    for weak, strong in WEAK_VERBS.items():
        if new.lower().startswith(weak):
            new = strong + new[len(weak):]; break
        if weak in new.lower():
            new = re.sub(weak, strong.lower(), new, count=1, flags=re.I)
    hit = [k for k in kws if k in new.lower()]
    return new, hit

def build_tailored(cv, jd, target_title):
    kws = jd_keywords(jd)
    out = {"title": target_title, "keywords": kws}
    skills_lines = (cv["sections"].get("skills",[]) +
                    cv["sections"].get("technical skills",[]) +
                    cv["sections"].get("soft skills",[]))
    skills = []
    for l in skills_lines:
        for part in re.split(r"[,\u2022•|]", l):
            p = part.strip()
            if 2 < len(p) < 45: skills.append(p)
    ranked = sorted(set(skills),
                    key=lambda s: -sum(1 for k in kws if k in s.lower()))
    out["skills"] = ranked

    exp_lines = cv["sections"].get("experience", [])
    bullets = extract_bullets(exp_lines)
    tailored = []
    for b in bullets:
        nb, hits = tailor_bullet(b, kws)
        tailored.append({"text": nb, "hits": hits,
                         "score": len(hits) + (1 if re.search(r"\d", nb) else 0)})
    tailored.sort(key=lambda x: -x["score"])
    out["bullets"] = tailored
    out["gaps"] = [k for k in kws
                   if k not in " ".join(skills + bullets).lower()][:12]
    out["summary"] = (
        f"{cv.get('headline') or 'Graduate'} seeking a {target_title} role. "
        f"Hands-on with {', '.join(ranked[:3]) if ranked else 'core tools'}. "
        f"[Add one measurable achievement here.]")
    out["education"] = cv["sections"].get("education", [])
    out["certs"] = cv["sections"].get("certification", []) + \
                   cv["sections"].get("certifications", [])
    out["experience_raw"] = exp_lines
    return out


def make_docx(cv, t):
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    d = docx.Document()
    for s in d.sections:
        s.top_margin = s.bottom_margin = Pt(36)
        s.left_margin = s.right_margin = Pt(46)
    n = d.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(10.5)

    h = d.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(cv["name"]); r.bold = True; r.font.size = Pt(19)
    sub = d.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(t["title"]).font.size = Pt(11.5)
    c = d.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.add_run(cv["contact"]).font.size = Pt(9.5)

    def head(txt):
        p = d.add_paragraph(); p.space_before = Pt(11); p.space_after = Pt(2)
        rr = p.add_run(txt.upper()); rr.bold = True; rr.font.size = Pt(11)
        rr.font.color.rgb = RGBColor(0x1F,0x2A,0x44)

    head("Professional Summary"); d.add_paragraph(t["summary"])
    head("Education")
    for l in t["education"]: d.add_paragraph(l)
    head("Skills")
    if t["skills"]: d.add_paragraph(" • ".join(t["skills"][:16]))
    head("Experience")
    for l in t["experience_raw"]:
        if not re.match(r"^[-•·▪*]", l.strip()): d.add_paragraph(l)
    for b in t["bullets"]:
        d.add_paragraph(b["text"], style="List Bullet")
    if t["certs"]:
        head("Certifications")
        for l in t["certs"]: d.add_paragraph(l)

    buf = io.BytesIO(); d.save(buf); buf.seek(0)
    return buf.getvalue()


def make_pdf(cv, t):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph,
                                    ListFlowable, ListItem, HRFlowable)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=14*mm,
                            bottomMargin=14*mm, leftMargin=16*mm, rightMargin=16*mm)
    ss = getSampleStyleSheet()
    nm = ParagraphStyle("nm", parent=ss["Title"], fontSize=19, spaceAfter=2)
    rl = ParagraphStyle("rl", parent=ss["Normal"], fontSize=11.5,
                        alignment=1, spaceAfter=2)
    ct = ParagraphStyle("ct", parent=ss["Normal"], fontSize=9,
                        alignment=1, textColor=colors.grey, spaceAfter=8)
    hd = ParagraphStyle("hd", parent=ss["Heading2"], fontSize=11,
                        spaceBefore=9, spaceAfter=3,
                        textColor=colors.HexColor("#1F2A44"))
    bd = ParagraphStyle("bd", parent=ss["Normal"], fontSize=10, leading=13.5)

    def esc(s): return (s or "").replace("&","&amp;").replace("<","&lt;")

    f = [Paragraph(esc(cv["name"]), nm), Paragraph(esc(t["title"]), rl),
         Paragraph(esc(cv["contact"]), ct),
         HRFlowable(width="100%", color=colors.HexColor("#CBD5E1")),
         Paragraph("PROFESSIONAL SUMMARY", hd), Paragraph(esc(t["summary"]), bd)]

    f.append(Paragraph("EDUCATION", hd))
    for l in t["education"]: f.append(Paragraph(esc(l), bd))
    if t["skills"]:
        f.append(Paragraph("SKILLS", hd))
        f.append(Paragraph(esc(" • ".join(t["skills"][:16])), bd))
    f.append(Paragraph("EXPERIENCE", hd))
    for l in t["experience_raw"]:
        if not re.match(r"^[-•·▪*]", l.strip()):
            f.append(Paragraph(f"<b>{esc(l)}</b>", bd))
    if t["bullets"]:
        f.append(ListFlowable(
            [ListItem(Paragraph(esc(b["text"]), bd), leftIndent=10)
             for b in t["bullets"]], bulletType="bullet", start="•", leftIndent=12))
    if t["certs"]:
        f.append(Paragraph("CERTIFICATIONS", hd))
        for l in t["certs"]: f.append(Paragraph(esc(l), bd))

    doc.build(f); buf.seek(0)
    return buf.getvalue()


# ═══════════════════════════════════════ sidebar
with st.sidebar:
    st.markdown("### 📍 where")
    all_saudi = st.checkbox("All of Saudi Arabia", value=True)
    chosen = [] if all_saudi else st.multiselect("Cities", list(SAUDI_CITIES),
                                                 default=["Riyadh","Medina"])
    keep_blank = st.checkbox("Include unlisted locations", value=False)

    st.markdown("### 🎓 level")
    level = st.radio("Show", ["Entry level", "Experienced"], index=0)
    entry_only = level == "Entry level"

    st.markdown("### 🔌 sources")
    use_private = st.checkbox("Startup ATS boards (fast)", value=True)
    use_gov = st.checkbox(f"Gov portals ({len(GOV_ENTITIES)})", value=True)
    use_emp = st.checkbox(f"Top employers ({len(EMPLOYERS)})", value=False,
                          help="Adds ~7 min on a cold run. Cached 12h after.")
    use_browser = st.checkbox("Real browser (renders JS)", value=True)
    use_boards = st.checkbox("Indeed + Bayt", value=True,
                             help="LinkedIn excluded — it blocks hard.")

    st.markdown("### 🧠 matching")
    use_semantic = st.checkbox("Semantic re-rank", value=True,
        help="Reads meaning, not just keywords. Runs locally, ~10s.")
    only_new = st.checkbox("Only show new since last run", value=False)
    min_score = st.slider("Min match score", 0, 200, 20)

    st.markdown("---")
    st.caption(f"Polite mode · {MIN_DELAY}–{MAX_DELAY}s delays · "
               f"{CACHE_HOURS}h cache · auto back-off on 429")


# ═══════════════════════════════════════ game
GAME_HTML = """
<div style="font-family:system-ui;text-align:center;padding:2px">
 <div style="display:flex;justify-content:center;gap:14px;margin-bottom:6px;
      font-weight:600;font-size:.8rem;color:#4B5563">
   <span id="lvl">lvl 1/5</span><span id="mv">moves 0</span><span id="stars"></span>
 </div>
 <div id="tubes" style="display:flex;gap:9px;justify-content:center;
      align-items:flex-end;height:146px"></div>
 <div id="msg" style="margin-top:6px;font-size:.95rem;font-weight:800;
      min-height:20px"></div>
 <button id="rst" style="margin-top:2px;padding:5px 14px;border-radius:10px;
   border:2px solid #111;background:#FDE68A;font-weight:700;cursor:pointer;
   font-size:.76rem;box-shadow:2px 2px 0 #111">restart</button>
</div>
<style>
.tube{width:33px;border:3px solid #111;border-top:none;border-radius:0 0 17px 17px;
 display:flex;flex-direction:column-reverse;align-items:center;padding-bottom:4px;
 gap:2px;cursor:pointer;background:#fff;
 transition:transform .18s cubic-bezier(.34,1.56,.64,1)}
.tube.sel{transform:translateY(-10px)}
.ball{width:23px;height:21px;border-radius:50%;border:2px solid rgba(0,0,0,.22);
 animation:drop .22s cubic-bezier(.34,1.56,.64,1)}
@keyframes drop{from{transform:translateY(-16px);opacity:.4}
to{transform:translateY(0);opacity:1}}
</style>
<script>
const C=['#EF4444','#3B82F6','#22C55E','#EAB308','#A855F7','#EC4899'];
const L=[{c:2,e:1,h:3},{c:3,e:1,h:3},{c:4,e:2,h:4},{c:5,e:2,h:4},{c:6,e:2,h:4}];
let lvl=0,tubes=[],sel=null,moves=0,stars=0,cap=4,lock=false;
const $=id=>document.getElementById(id);

function build(){
  const s=L[lvl];cap=s.h;let b=[];
  for(let c=0;c<s.c;c++)for(let k=0;k<s.h;k++)b.push(c);
  for(let i=b.length-1;i>0;i--){const j=0|Math.random()*(i+1);[b[i],b[j]]=[b[j],b[i]];}
  tubes=[];
  for(let t=0;t<s.c;t++)tubes.push(b.slice(t*s.h,(t+1)*s.h));
  for(let e=0;e<s.e;e++)tubes.push([]);
  moves=0;sel=null;lock=false;$('msg').textContent='';draw();
}
const won=()=>tubes.every(t=>!t.length||(t.length===cap&&t.every(x=>x===t[0])));

function draw(){
  const box=$('tubes');box.innerHTML='';
  $('lvl').textContent='lvl '+(lvl+1)+'/5';
  $('mv').textContent='moves '+moves;
  $('stars').textContent='⭐'.repeat(stars);
  tubes.forEach((t,i)=>{
    const d=document.createElement('div');
    d.className='tube'+(sel===i?' sel':'');
    d.style.height=(cap*25+8)+'px';
    t.forEach(c=>{const e=document.createElement('div');
      e.className='ball';e.style.background=C[c];d.appendChild(e);});
    d.onclick=()=>{
      if(lock)return;
      if(sel===null){if(t.length)sel=i;}
      else if(sel===i){sel=null;}
      else{
        const f=tubes[sel],o=tubes[i];
        if(f.length&&o.length<cap&&(!o.length||o[o.length-1]===f[f.length-1])){
          o.push(f.pop());moves++;
        }
        sel=null;
        if(won()){
          lock=true;
          if(lvl<4){
            $('msg').innerHTML='<span style="color:#059669">clear! 🫧</span>';
            stars++;lvl++;setTimeout(build,750);draw();return;
          }
          stars=5;
          $('msg').innerHTML='<div style="font-size:1.2rem">⭐⭐⭐⭐⭐</div>'+
            '<div style="color:#DB2777;font-size:1.02rem">GOOD GIRL LULUTY!</div>';
          draw();return;
        }
      }
      draw();
    };
    box.appendChild(d);
  });
}
$('rst').onclick=build;build();
</script>
"""

# ═══════════════════════════════════════ tabs
t_search, t_cv, t_track = st.tabs(["🔎 Find jobs", "📝 CV lab", "📋 Applications"])
profile, cv_text, field, field_kws, label = None, "", None, [], ""

with t_search:
    left, right = st.columns([5, 3])

    with left:
        m1, m2, m3 = st.tabs(["Upload CV", "Paste JD", "Keywords"])
        with m1:
            up = st.file_uploader("Drop the CV", type=["pdf","docx","txt"],
                                  key="cvsearch")
            if up:
                cv_text = read_upload(up)
                st.session_state["cv_text"] = cv_text
                profile, label = build_cv_profile(cv_text), "cv"
                field, field_kws = detect_field(cv_text)
                fr = _fused_ratio(cv_text)
                st.success(f"Parsed {len(cv_text):,} chars")
                if fr > 0.02:
                    st.warning("Some words look fused together in this PDF. "
                               "Matching may be weaker — a .docx version works better.")
                if field:
                    st.info(f"Field: **{field.title()}**")
                st.markdown("".join(f'<span class="badge b-blue">{k}</span>'
                            for k in list(profile)[:12]), unsafe_allow_html=True)
        with m2:
            jd = st.text_area("Paste the job description", height=170)
            if jd.strip():
                profile, label = build_profile(jd), "jd"
                st.session_state["jd_text"] = jd
        with m3:
            kw = st.text_input("Titles / keywords",
                "data analyst, business analyst, information systems, database")
            if kw.strip() and not profile:
                profile = {t.strip().lower():5 for t in kw.split(",") if t.strip()}
                label = "keywords"

        go = st.button("🐸  LET'S GOOO", type="primary", use_container_width=True)

    with right:
        st.markdown("##### 🫧 ball sort — play while it scans")
        components.html(GAME_HTML, height=250)

    if go:
        if not profile:
            st.error("Give me a CV, a JD, or keywords first."); st.stop()
        if not all_saudi and not chosen:
            st.error("Pick a city."); st.stop()

        employers_ats = _load(CACHE, {})
        jobcache = load_cache()
        results, dead, notes = [], [], []
        bar, status = st.progress(0.0), st.empty()

        if use_private:
            if not employers_ats:
                status.info("First run — probing ATS APIs. Go sort some balls.")
                for i,slug in enumerate(CANDIDATE_SLUGS):
                    a = probe(slug)
                    if a: employers_ats[slug]=a
                    bar.progress((i+1)/len(CANDIDATE_SLUGS)*0.25)
                    polite_sleep()
                _save(CACHE, employers_ats)
            items = list(employers_ats.items())
            for i,(slug,ats) in enumerate(items):
                k=f"ats:{slug}"
                if k in jobcache and cache_fresh(jobcache[k]):
                    jobs=jobcache[k]["jobs"]
                else:
                    status.write(f"🌐 {slug}…")
                    jobs=fetch_ats(slug,ats)
                    jobcache[k]={"ts":datetime.now().isoformat(),"jobs":jobs}
                    polite_sleep()
                for j in jobs:
                    if loc_ok(j.get("location"),chosen,all_saudi,keep_blank):
                        p,h,g=score(j,profile,entry_only,field_kws)
                        if p>=min_score:
                            j["score"],j["matched"],j["gaps"],j["src"]=p,h,g,"ATS"
                            results.append(j)
                bar.progress(0.25+(i+1)/max(len(items),1)*0.1)

        targets = []
        if use_gov: targets += [("GOV",n,u) for n,u in GOV_ENTITIES.items()]
        if use_emp: targets += [("EMP",n,u) for n,u in EMPLOYERS.items()]

        if targets:
            todo, cached_hits = [], []
            for src,name,url in targets:
                k=f"site:{name}"
                if k in jobcache and cache_fresh(jobcache[k]):
                    for j in jobcache[k]["jobs"]:
                        j["src"]=src; cached_hits.append(j)
                else:
                    todo.append((name,url,src))

            fresh = {}
            if todo:
                srcmap = {n:s for n,_,s in todo}
                pairs = [(n,u) for n,u,_ in todo]
                if use_browser:
                    status.write(f"🌐 Opening {len(pairs)} career portals…")
                    def _prog(i,n,nm):
                        status.write(f"🏛 {nm}  ({i}/{n})")
                        bar.progress(0.35 + i/n*0.5)
                    try:
                        fresh, d2 = fetch_site_browser(pairs, _prog)
                        dead.extend(d2)
                    except ImportError:
                        notes.append("Playwright missing — pip install playwright "
                                     "&& python -m playwright install chromium")
                        use_browser = False
                    except Exception as e:
                        notes.append(f"Browser mode failed ({e}); used plain fetch.")
                        use_browser = False

                if not use_browser:
                    for i,(n,u) in enumerate(pairs):
                        status.write(f"🏛 {n}…")
                        got = fetch_site_plain(n,u)
                        if got: fresh[n]=got
                        else: dead.append(n)
                        bar.progress(0.35+(i+1)/len(pairs)*0.5)
                        polite_sleep()

                for n,jl in fresh.items():
                    for j in jl: j["src"] = srcmap.get(n,"GOV")
                    jobcache[f"site:{n}"]={"ts":datetime.now().isoformat(),"jobs":jl}
                for n,_ in pairs:
                    if n not in fresh:
                        jobcache[f"site:{n}"]={"ts":datetime.now().isoformat(),"jobs":[]}

            for j in cached_hits + [x for jl in fresh.values() for x in jl]:
                p,h,g = score(j, profile, entry_only, field_kws)
                if p >= min_score:
                    j["score"], j["matched"], j["gaps"] = p, h, g
                    results.append(j)
            bar.progress(0.87)

        if use_boards:
            status.write("🔍 Indeed + Bayt…")
            jobs,err = fetch_jobspy(list(profile)[0], chosen or ["Riyadh"], all_saudi)
            if err: notes.append(err)
            for j in jobs:
                if loc_ok(j.get("location"),chosen,all_saudi,True):
                    p,h,g=score(j,profile,entry_only,field_kws)
                    if p>=min_score:
                        j["score"],j["matched"],j["gaps"],j["src"]=p,h,g,"BOARD"
                        results.append(j)
            bar.progress(0.94)

        save_cache(jobcache)

        seen,uniq=set(),[]
        for j in sorted(results,key=lambda x:-x["score"]):
            key=(j["title"].lower(), j.get("company",""))
            if key not in seen: seen.add(key); uniq.append(j)

        if use_semantic and label == "cv" and st.session_state.get("cv_text"):
            try:
                status.write("🧠 Re-ranking by meaning…")
                uniq = semantic_rerank(uniq, st.session_state["cv_text"])
            except Exception as e:
                notes.append(f"Semantic re-rank skipped: {e}")

        seen_db = _load(SEENFILE, {})
        now_iso = datetime.now().isoformat()
        n_new = 0
        for j in uniq:
            k = job_key(j)
            if k in seen_db:
                j["is_new"] = False; j["first_seen"] = seen_db[k]
            else:
                j["is_new"] = True; j["first_seen"] = now_iso; n_new += 1
            seen_db[k] = j["first_seen"]
        _save(SEENFILE, seen_db)

        bar.progress(1.0); status.empty()
        shown = [j for j in uniq if j["is_new"]] if only_new else uniq

        c1,c2,c3 = st.columns(3)
        c1.metric("Matches", len(uniq))
        c2.metric("New since last run", n_new)
        c3.metric("Showing", len(shown))
        st.markdown(f"#### {random.choice(HYPE)}")

        if not shown:
            st.warning("Nothing to show. Lower min score, widen cities, "
                       "untick 'only new', or try 'Experienced'.")

        for j in shown:
            bd={"ATS":"b-green","GOV":"b-blue","EMP":"b-amber",
                "BOARD":"b-pink"}.get(j.get("src"),"b-gray")
            newbadge = '<span class="badge b-new">NEW</span>' if j.get("is_new") else ''
            posted = (f'<span class="badge b-gray">📅 {j["posted"]}</span>'
                      if j.get("posted") else '')
            chips = "".join(f'<span class="badge b-gray">{m}</span>'
                            for m in j.get("matched",[]))
            gapline = ""
            if j.get("gaps"):
                total = len(j.get("matched",[])) + len(j["gaps"])
                gapline = ('<div style="margin-top:8px;font-size:.8rem;color:#6B7280">'
                           f'Matches {len(j.get("matched",[]))} of {total} key skills'
                           ' · missing: ' +
                           " ".join(f'<span class="badge b-red">{g}</span>'
                                    for g in j["gaps"]) + '</div>')
            fit = int(j.get("sim",0)*100)
            fitbar = (f'<div class="fitbar"><div class="fitfill" style="width:{fit}%">'
                      f'</div></div><div style="font-size:.7rem;color:#9CA3AF;'
                      f'margin-top:3px">{fit}% semantic fit</div>') if j.get("sim") else ''

            card = (
f'<div class="jobcard">'
f'<div style="display:flex;justify-content:space-between;gap:12px">'
f'<div style="flex:1">'
f'<div style="font-size:1.08rem;font-weight:700;line-height:1.3;color:#111827">'
f'{j["title"]} {newbadge}</div>'
f'<div style="color:#6B7280;font-size:.84rem;margin:4px 0 8px">'
f'{j.get("company","")} · {j.get("location") or "location n/a"} '
f'<span class="badge {bd}">{j.get("src","")}</span> {posted}</div>'
f'{chips}{gapline}</div>'
f'<div style="text-align:right;min-width:66px">'
f'<div style="font-size:1.5rem;font-weight:700;color:#A855F7">{j["score"]}</div>'
f'<div style="font-size:.66rem;color:#9CA3AF">score</div></div>'
f'</div>{fitbar}'
f'<div style="margin-top:10px"><a href="{j["url"]}" target="_blank" '
f'style="font-weight:600;color:#3B82F6;text-decoration:none">open posting ↗</a></div>'
f'</div>'
            )
            st.markdown(card, unsafe_allow_html=True)

        if shown:
            import csv as _csv
            b=io.StringIO(); w=_csv.writer(b)
            w.writerow(["score","fit%","new","source","company","title",
                        "location","posted","matched","gaps","url"])
            for j in shown:
                w.writerow([j["score"], int(j.get("sim",0)*100),
                            "YES" if j.get("is_new") else "",
                            j.get("src",""), j.get("company",""), j["title"],
                            j.get("location",""), j.get("posted",""),
                            ", ".join(j.get("matched",[])),
                            ", ".join(j.get("gaps",[])), j["url"]])
            st.download_button("⬇ Download CSV", b.getvalue(),
                f"luluty_{datetime.now():%Y-%m-%d}.csv","text/csv")

            with st.expander("➕ Add jobs to the application tracker"):
                picks = st.multiselect("Which ones did you apply to?",
                    [f"{j['title']} — {j.get('company','')}" for j in shown[:40]])
                if st.button("Save to tracker") and picks:
                    apps = _load(TRACKFILE, [])
                    have = {(a["title"], a["company"]) for a in apps}
                    for j in shown:
                        lbl = f"{j['title']} — {j.get('company','')}"
                        if lbl in picks and (j["title"], j.get("company","")) not in have:
                            apps.append({"title":j["title"],
                                         "company":j.get("company",""),
                                         "url":j["url"], "status":"Applied",
                                         "date":datetime.now().strftime("%Y-%m-%d"),
                                         "notes":""})
                    _save(TRACKFILE, apps)
                    st.success(f"Added {len(picks)}. See the Applications tab.")

        if dead:
            with st.expander(f"⚠ {len(dead)} portals returned nothing"):
                st.write(", ".join(dead))
                st.caption("No openings, JS the scraper can't reach, or a login "
                           "wall. Worth opening manually.")
        for n in notes: st.warning(n)


# ═══════════════════════════════════════ CV LAB
with t_cv:
    st.markdown("### 📝 CV lab — tailor the CV to one job")
    st.caption("Reorders skills, strengthens verbs, surfaces the JD's language, "
               "and exports a single-column ATS-safe file. It will not invent "
               "experience — fabrications fail at interview and employers verify.")

    c1, c2 = st.columns(2)
    with c1:
        u2 = st.file_uploader("CV", type=["pdf","docx","txt"], key="cvlab")
        if u2:
            st.session_state["cv_text"] = read_upload(u2)
        text = st.session_state.get("cv_text","")
        if text:
            st.success(f"Loaded — {len(text):,} chars")
            if _fused_ratio(text) > 0.02:
                st.warning("Words look fused in this PDF — a .docx works better.")
    with c2:
        target_jd = st.text_area("Paste the job description you're applying to",
                                 height=190, key="tjd")

    if text and target_jd.strip():
        cv = parse_cv(text)
        title = st.text_input("Target job title (appears under the name)",
                              jd_title(target_jd))
        t = build_tailored(cv, target_jd, title)

        st.markdown("#### Skills — JD matches first")
        st.markdown("".join(
            f'<span class="badge {"b-green" if any(k in s.lower() for k in t["keywords"]) else "b-gray"}">{s}</span>'
            for s in t["skills"][:20]), unsafe_allow_html=True)

        st.markdown("#### Experience bullets — strongest match first")
        for b in t["bullets"]:
            tag = (f'<span class="badge b-amber">{len(b["hits"])} JD hits</span>'
                   if b["hits"] else '<span class="badge b-gray">no match</span>')
            st.markdown(f'<div style="margin-bottom:7px">{tag} {b["text"]}</div>',
                        unsafe_allow_html=True)

        if t["gaps"]:
            st.markdown("#### ⚠ JD terms missing from the CV")
            st.caption("Add only what's genuinely true.")
            st.markdown("".join(f'<span class="badge b-pink">{g}</span>'
                        for g in t["gaps"]), unsafe_allow_html=True)

        st.markdown("#### Summary line — edit before sending")
        t["summary"] = st.text_area("Summary", t["summary"], height=80)

        st.markdown("#### 📥 Download")
        d1, d2 = st.columns(2)
        safe = re.sub(r"\W+","_", title)[:30]
        try:
            d1.download_button("⬇ Word (.docx)", make_docx(cv, t), f"CV_{safe}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True)
        except Exception as e:
            d1.error(f"DOCX failed: {e}")
        try:
            d2.download_button("⬇ PDF", make_pdf(cv, t), f"CV_{safe}.pdf",
                "application/pdf", use_container_width=True)
        except Exception as e:
            d2.error(f"PDF failed — pip install reportlab ({e})")

        with st.expander("Before sending"):
            st.markdown("""
- Read every bullet. If a rephrase changed the meaning, fix it.
- Add real numbers where gaps were flagged — how many queries, events, reports.
- Match the job title exactly as the posting writes it.
- Send **.docx** unless the portal demands PDF.
- One page for entry level.
            """)
    else:
        st.info("Load a CV and paste a JD to generate a tailored version.")


# ═══════════════════════════════════════ TRACKER
with t_track:
    st.markdown("### 📋 Applications")
    apps = _load(TRACKFILE, [])

    if not apps:
        st.info("Nothing tracked yet. Run a search, then use "
                "'Add jobs to the application tracker' under the results.")
    else:
        counts = Counter(a.get("status","Applied") for a in apps)
        cols = st.columns(5)
        for c,(lbl,key) in zip(cols, [("Total",None),("Applied","Applied"),
                ("Interview","Interview"),("Offer","Offer"),("Rejected","Rejected")]):
            c.metric(lbl, len(apps) if key is None else counts.get(key,0))

        st.markdown("")
        STATUSES = ["Applied","Follow up","Interview","Offer","Rejected","Closed"]
        changed = False
        for i,a in enumerate(apps):
            with st.container(border=True):
                c1,c2,c3 = st.columns([4,2,1])
                with c1:
                    st.markdown(f"**{a['title']}**")
                    st.caption(f"{a['company']} · applied {a['date']}")
                    if a.get("url"): st.link_button("Open posting", a["url"])
                with c2:
                    cur = a.get("status","Applied")
                    ns = st.selectbox("Status", STATUSES,
                        index=STATUSES.index(cur) if cur in STATUSES else 0,
                        key=f"st{i}", label_visibility="collapsed")
                    if ns != cur: a["status"]=ns; changed=True
                    nn = st.text_input("Notes", a.get("notes",""),
                        key=f"nt{i}", placeholder="notes…",
                        label_visibility="collapsed")
                    if nn != a.get("notes",""): a["notes"]=nn; changed=True
                with c3:
                    if st.button("🗑", key=f"dl{i}"):
                        apps.pop(i); _save(TRACKFILE, apps); st.rerun()
        if changed:
            _save(TRACKFILE, apps)

        st.markdown("")
        import csv as _csv2
        b=io.StringIO(); w=_csv2.writer(b)
        w.writerow(["date","company","title","status","notes","url"])
        for a in apps:
            w.writerow([a["date"],a["company"],a["title"],
                        a.get("status",""),a.get("notes",""),a.get("url","")])
        st.download_button("⬇ Export tracker", b.getvalue(),
                           "applications.csv","text/csv")
